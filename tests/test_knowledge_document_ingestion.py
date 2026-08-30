"""Markdown / PDF / DOCX 统一 Knowledge Document Contract 静态测试。"""

from __future__ import annotations

from pathlib import Path

from agent.knowledge.chunking import KnowledgeChunker
from agent.knowledge.corpus import KnowledgeCorpus
from agent.knowledge.document_ingestion import DocxDocumentParser, DocumentIngestionError, PdfDocumentParser
from agent.knowledge.models import KnowledgeDocument

ROOT = Path(__file__).resolve().parents[1]


class _FakePdfPage:
    """模拟 pypdf PageObject，只暴露当前 Parser 实际依赖的 extract_text。"""

    def __init__(self, text: str):
        self.text = text

    def extract_text(self):
        """返回固定页面文本。"""
        return self.text


class _FakePdfReader:
    """模拟多页 PdfReader。"""

    def __init__(self, pages):
        self.pages = pages


def test_current_manifest_remains_18_governed_markdown_documents():
    docs = KnowledgeCorpus(ROOT).load()
    assert len(docs) == 18
    assert {doc.source_format for doc in docs} == {"markdown"}


def test_text_pdf_is_normalized_to_blocks_with_page_provenance(tmp_path):
    path = tmp_path / "runbook.pdf"
    path.write_bytes(b"%PDF-static-test")
    parser = PdfDocumentParser(
        reader_factory=lambda _: _FakePdfReader(
            [_FakePdfPage("Failure recovery\n\nCheck exact partition."), _FakePdfPage("Escalate if incomplete.")]
        )
    )
    parsed = parser.parse(path)
    assert parsed.source_format == "pdf"
    assert [block.page_number for block in parsed.blocks] == [1, 1, 2]
    assert all(block.block_type == "paragraph" for block in parsed.blocks)

    document = KnowledgeDocument(
        document_id="test.pdf",
        title="PDF Runbook",
        scope="runbook",
        domain="commerce",
        authority="runbook",
        owner="data-platform",
        status="active",
        tags=(),
        source_path=Path("knowledge/runbooks/test.pdf"),
        reviewed_at=None,
        body=parsed.body,
        document_sha256="a" * 64,
        source_format=parsed.source_format,
        blocks=parsed.blocks,
    )
    chunks = KnowledgeChunker(target_chars=400, max_chars=2400).chunk_document(document)
    assert chunks
    assert chunks[0].source_format == "pdf"
    assert chunks[0].page_numbers == (1,)


def test_scanned_pdf_without_text_fails_closed(tmp_path):
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-static-test")
    parser = PdfDocumentParser(reader_factory=lambda _: _FakePdfReader([_FakePdfPage(""), _FakePdfPage(" ")]))
    try:
        parser.parse(path)
    except DocumentIngestionError as exc:
        assert "OCR" in str(exc)
    else:
        raise AssertionError("image/scanned PDF must not be treated as parsed text")


def test_docx_heading_paragraph_and_table_are_normalized(tmp_path):
    from docx import Document

    path = tmp_path / "policy.docx"
    doc = Document()
    doc.add_heading("Recovery", level=1)
    doc.add_paragraph("Check the exact partition before replay.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Condition"
    table.cell(0, 1).text = "Action"
    doc.save(path)

    parsed = DocxDocumentParser().parse(path)
    assert parsed.source_format == "docx"
    assert [block.block_type for block in parsed.blocks] == ["heading", "paragraph", "table"]
    assert parsed.blocks[1].section == "Recovery"


def test_docx_can_enter_governed_corpus_when_manifest_supplies_binary_metadata(tmp_path):
    """DOCX 没有 Front Matter 时，Manifest 的完整治理元数据可以驱动统一 Corpus Contract。"""
    import yaml
    from docx import Document

    (tmp_path / "knowledge/runbooks").mkdir(parents=True)
    (tmp_path / "metadata/knowledge").mkdir(parents=True)
    path = tmp_path / "knowledge/runbooks/recovery.docx"
    doc = Document()
    doc.add_heading("Recovery", level=1)
    doc.add_paragraph("Check exact partition completeness before replay.")
    doc.save(path)

    manifest = {
        "version": 2,
        "corpus": "test",
        "status": "active",
        "documents": [
            {
                "id": "test.recovery.docx",
                "path": "knowledge/runbooks/recovery.docx",
                "title": "Recovery DOCX",
                "scope": "runbook",
                "domain": "commerce",
                "authority": "runbook",
                "owner": "data-platform",
                "status": "active",
                "tags": ["recovery"],
            }
        ],
    }
    (tmp_path / "metadata/knowledge/corpus_manifest.yml").write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8"
    )
    docs = KnowledgeCorpus(tmp_path).load()
    assert len(docs) == 1
    assert docs[0].source_format == "docx"
    assert docs[0].blocks and docs[0].blocks[0].block_type == "heading"
    chunks = KnowledgeChunker().chunk_documents(docs)
    assert chunks and chunks[0].source_format == "docx"
