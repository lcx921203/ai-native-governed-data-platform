"""多格式 Knowledge Document Ingestion（知识文档接入）适配层。

业务逻辑：Markdown / PDF / DOCX 先解析成统一 ``ParsedKnowledgeDocument``，再由
``KnowledgeCorpus`` 施加 Manifest / Owner / Scope / Authority 治理，最后进入同一 Chunker。

API 知识：
- PDF 使用 ``pypdf.PdfReader`` 读取已有文本层；
- DOCX 使用 ``python-docx`` 并按 Word XML 原始顺序遍历 Paragraph / Table；
- Markdown 仍保留 YAML Front Matter，由 Corpus 层做治理校验。

工程边界：当前不做 OCR。扫描 PDF 如果没有可提取文本会 Fail Closed，并明确要求
后续 OCR / layout pipeline；不会把空文本当成成功 ingestion。
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable

from agent.knowledge.models import KnowledgeBlock


class DocumentIngestionError(RuntimeError):
    """文档格式不支持、解析器缺失或内容无法可靠抽取时抛出的契约错误。"""


@dataclass(frozen=True)
class ParsedKnowledgeDocument:
    """文件格式解析后的统一中间结果。

    ``body`` 是供兼容路径和调试使用的规范化文本；``blocks`` 才是 PDF / DOCX
    进入 Structure-aware Chunker 的结构化输入。
    """

    source_format: str
    body: str
    blocks: tuple[KnowledgeBlock, ...]
    parser: str


class MarkdownDocumentParser:
    """Markdown 文本读取器；Front Matter 治理由 ``KnowledgeCorpus`` 负责。"""

    source_format = "markdown"

    def parse(self, path: Path) -> ParsedKnowledgeDocument:
        """按 UTF-8 读取 Markdown 原文。

        这里不提前切片，也不解释 Front Matter；保持已有 Markdown Chunk identity 稳定，
        后续仍由 Corpus + Markdown-aware Chunker 完成治理和结构感知切片。
        """
        raw = path.read_text(encoding="utf-8")
        return ParsedKnowledgeDocument(self.source_format, raw, (), "markdown-utf8")


class PdfDocumentParser:
    """文本型 PDF 解析器；保留 page provenance，但不承担 OCR。"""

    source_format = "pdf"

    def __init__(self, *, reader_factory: Callable[[Path], Any] | None = None):
        """允许测试注入 Fake Reader；生产路径按需导入 ``pypdf.PdfReader``。"""
        self.reader_factory = reader_factory

    def _reader(self, path: Path) -> Any:
        """构造 PDF Reader；缺少可选依赖时给出明确安装提示。"""
        if self.reader_factory is not None:
            return self.reader_factory(path)
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - 由 requirements-rag.txt 负责运行依赖
            raise DocumentIngestionError("pypdf 未安装；请安装 requirements-rag.txt") from exc
        return PdfReader(str(path))

    def parse(self, path: Path) -> ParsedKnowledgeDocument:
        """提取 PDF 已存在的文本层，并按页保留段落结构。

        每页通过 ``extract_text()`` 获取文本；空页会跳过。如果整份 PDF 没有任何文本，
        视为扫描件/图片型 PDF，当前源码直接拒绝，避免把 OCR 缺失伪装成成功解析。
        """
        reader = self._reader(path)
        blocks: list[KnowledgeBlock] = []
        page_bodies: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            page_bodies.append(text)
            paragraphs = [part.strip() for part in text.split("\n\n") if part.strip()]
            if not paragraphs:
                paragraphs = [text]
            for paragraph in paragraphs:
                blocks.append(
                    KnowledgeBlock(
                        block_type="paragraph",
                        text=paragraph,
                        section=f"Page {page_number}",
                        page_number=page_number,
                    )
                )
        if not blocks:
            raise DocumentIngestionError(
                "PDF 没有可提取文本层；当前版本不执行 OCR。请先走 OCR/layout ingestion，再进入 Knowledge Corpus。"
            )
        return ParsedKnowledgeDocument(self.source_format, "\n\n".join(page_bodies), tuple(blocks), "pypdf")


class DocxDocumentParser:
    """DOCX 解析器；按 Word 文档顺序保留标题、段落和表格。"""

    source_format = "docx"

    @staticmethod
    def _iter_blocks(document: Any) -> Iterable[Any]:
        """按底层 Word XML 顺序迭代 Paragraph / Table。

        ``python-docx`` 的 ``document.paragraphs`` 与 ``document.tables`` 分开访问会丢失两者交错顺序，
        因此这里直接遍历 ``document.element.body``，再包装成 Paragraph / Table 对象。
        """
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        for child in document.element.body.iterchildren():
            tag = str(child.tag)
            if tag.endswith("}p"):
                yield Paragraph(child, document)
            elif tag.endswith("}tbl"):
                yield Table(child, document)

    def parse(self, path: Path) -> ParsedKnowledgeDocument:
        """把 DOCX 转成统一 heading / paragraph / table blocks。

        Heading 样式更新当前 section；普通段落继承最近标题。表格按行列转成稳定文本，
        让后续 Chunker 不需要理解 Word XML。
        """
        try:
            from docx import Document
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:  # pragma: no cover - 由 requirements-rag.txt 负责运行依赖
            raise DocumentIngestionError("python-docx 未安装；请安装 requirements-rag.txt") from exc

        document = Document(str(path))
        section_stack: list[str] = []
        current_section = "Document"
        blocks: list[KnowledgeBlock] = []

        for item in self._iter_blocks(document):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if not text:
                    continue
                style_name = (getattr(item.style, "name", "") or "").strip()
                if style_name.casefold().startswith("heading"):
                    digits = "".join(ch for ch in style_name if ch.isdigit())
                    level = int(digits or "1")
                    section_stack[:] = section_stack[: max(0, level - 1)]
                    while len(section_stack) < level - 1:
                        section_stack.append("")
                    section_stack.append(text)
                    current_section = " > ".join(part for part in section_stack if part)
                    blocks.append(KnowledgeBlock("heading", text, current_section))
                else:
                    blocks.append(KnowledgeBlock("paragraph", text, current_section))
            elif isinstance(item, Table):
                rows = []
                for row in item.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    rows.append(" | ".join(cells))
                table_text = "\n".join(row for row in rows if row.strip(" |"))
                if table_text:
                    blocks.append(KnowledgeBlock("table", table_text, current_section))

        if not blocks:
            raise DocumentIngestionError("DOCX 没有可索引的标题、段落或表格内容")
        body = "\n\n".join(block.text for block in blocks)
        return ParsedKnowledgeDocument(self.source_format, body, tuple(blocks), "python-docx")


class KnowledgeDocumentParserRegistry:
    """根据受治理文件后缀选择固定 Parser；Agent 不能自行传任意 Parser。"""

    def __init__(self):
        """注册当前明确支持的三种文档格式。"""
        self._parsers = {
            ".md": MarkdownDocumentParser(),
            ".markdown": MarkdownDocumentParser(),
            ".pdf": PdfDocumentParser(),
            ".docx": DocxDocumentParser(),
        }

    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        """返回受治理 Parser Registry 当前支持的文件后缀。"""
        return tuple(sorted(self._parsers))

    def parse(self, path: Path) -> ParsedKnowledgeDocument:
        """按文件后缀调用固定 Parser；未知格式 Fail Closed。"""
        suffix = path.suffix.casefold()
        parser = self._parsers.get(suffix)
        if parser is None:
            raise DocumentIngestionError(
                f"不支持的 Knowledge 文档格式: {suffix or '<none>'}; allowed={self.supported_suffixes}"
            )
        return parser.parse(path)
